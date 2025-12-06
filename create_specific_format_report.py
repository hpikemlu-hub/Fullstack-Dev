import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.ns import qn
import fitz  # PyMuPDF

def create_report_with_specific_format():
    # Create a new document
    doc = Document()
    
    # Add title (we'll use the default style for now)
    title = doc.add_heading('LAPORAN KEGIATAN LEMBUR', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Pengembangan Aplikasi Workload Management System')
    run.bold = True
    
    # Add date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run('Tanggal 8 - 9 dan 29 - 30 November 2025')
    date_run.bold = True
    
    # Add space
    doc.add_paragraph()
    
    # A. Pendahuluan
    heading_a = doc.add_heading('A. Pendahuluan', level=1)
    
    para_a = doc.add_paragraph()
    para_a.add_run(
        'Dalam rangka mendukung efisiensi dan efektivitas pelaksanaan tugas di Direktorat Hukum dan Perjanjian Sosial Budaya, '
        'Kementerian Luar Negeri, telah dilaksanakan kegiatan lembur dalam rangka pengembangan Aplikasi Workload Management System. '
        'Kegiatan ini dilakukan sebagai tindak lanjut dari surat perintah kerja yang dikeluarkan untuk penyelesaian tugas-tugas '
        'penting terkait manajemen beban kerja pegawai dan pelaporan kinerja. Aplikasi ini dirancang untuk memberikan solusi '
        'dalam pengelolaan beban kerja, penjadwalan, dan dokumentasi kinerja pegawai di lingkungan Direktorat Hukum dan Perjanjian Sosial Budaya.'
    )
    
    doc.add_paragraph()
    
    # B. Kegiatan yang dilaksanakan
    heading_b = doc.add_heading('B. Kegiatan yang dilaksanakan', level=1)
    
    para_b = doc.add_paragraph()
    para_b.add_run(
        'Kegiatan lembur yang dilaksanakan dalam periode tanggal 8 - 9 dan 29 - 30 November 2025 meliputi pengembangan '
        'dan perbaikan sistem aplikasi Workload Management berdasarkan surat perintah kerja sebagai berikut:'
    )
    
    # Process PDF files to extract information
    pdf_folder = '/home/hpsb/workload-app/Laporan/'
    pdf_files = [
        'Dok-23500-Surat Perintah Kerja Lembur Pembaruan System Aplikasi Administrasion Tools Direktorat Hukum dan Perjanjian Sosial Budaya pada tanggal 8 November 2025-signed (1).pdf',
        'Dok-34545-Surat Perintah Kerja Lembur Pembaruan System Aplikasi Administrasion Tools Direktorat Hukum dan Perjanjian Sosial Budaya pada tanggal 15 November 2025-signed (1).pdf',
        'Dok-65030-Surat Perintah Kerja Lembur Pembaruan System Aplikasi Administrasion Tools Direktorat Hukum dan Perjanjian Sosial Budaya pada tanggal 9 November 2025-signed (1).pdf',
        'Dok-85247-Surat Perintah Kerja Lembur Pembaruan System Aplikasi Administrasion Tools Direktorat Hukum dan Perjanjian Sosial Budaya pada tanggal 16 November 2025-signed (1).pdf'
    ]
    
    for pdf_file in pdf_files:
        pdf_path = os.path.join(pdf_folder, pdf_file)
        
        try:
            with fitz.open(pdf_path) as pdf:
                text = ""
                for page_num in range(min(2, len(pdf))):  # Read first 2 pages
                    page = pdf[page_num]
                    text += page.get_text()
            
            # Extract basic information (this is a simplified extraction)
            lines = text.split('\n')
            # Find the document number and date
            doc_num = pdf_file.split('-')[0]  # Extract from filename
            
            # Find the date from the filename
            date_part = pdf_file.split('tanggal ')[1].split('-')[0] if 'tanggal' in pdf_file else 'Tanggal tidak ditemukan'
            
            # Create list item for each work order
            list_item = doc.add_paragraph(style='List Bullet')
            list_item.add_run(f'Surat Perintah Kerja No. {doc_num} tanggal {date_part}: ')
            list_item.add_run('Pembaruan System Aplikasi Administrasion Tools Direktorat Hukum dan Perjanjian Sosial Budaya')
            
        except Exception as e:
            list_item = doc.add_paragraph(style='List Bullet')
            list_item.add_run(f'Gagal membaca file {pdf_file}: {str(e)}')
    
    doc.add_paragraph()
    
    # C. Hasil yang dicapai
    heading_c = doc.add_heading('C. Hasil yang dicapai dalam kegiatan lembur Direktorat HP Sosial Budaya pada tanggal 8 - 9 dan 29 - 30 November 2025 telah terselesaikan hal hal berikut:', level=1)
    
    # Results list
    results = [
        'Pengembangan modul dashboard untuk menampilkan ikhtisar beban kerja secara real-time',
        'Implementasi sistem manajemen tugas (workload) yang memungkinkan pegawai untuk melihat dan mengelola tugas-tugas mereka',
        'Pembuatan modul tim tugas (team tasks) untuk melihat distribusi beban kerja dalam tim',
        'Integrasi modul kalender untuk manajemen jadwal kegiatan dan perjalanan dinas',
        'Pengembangan modul laporan dan analitik untuk evaluasi kinerja',
        'Implementasi sistem otentikasi dan otorisasi pengguna berdasarkan peran (role-based access)',
        'Penyempurnaan antarmuka pengguna (UI/UX) agar lebih intuitif dan responsif',
        'Integrasi dengan sistem e-kinerja untuk keperluan dokumentasi kinerja pegawai'
    ]
    
    for result in results:
        list_item = doc.add_paragraph(style='List Bullet')
        list_item.add_run(result)
    
    doc.add_paragraph()
    
    # D. Kesimpulan dan Saran
    heading_d = doc.add_heading('D. Kesimpulan dan Saran', level=1)
    
    # Kesimpulan
    para_d1 = doc.add_paragraph()
    para_d1.add_run(
        'Kesimpulan: Kegiatan lembur dalam rangka pengembangan Aplikasi Workload Management System telah berhasil '
        'dilaksanakan sesuai dengan surat perintah kerja yang dikeluarkan. Aplikasi ini diharapkan dapat meningkatkan '
        'efisiensi dan transparansi dalam manajemen beban kerja di lingkungan Direktorat Hukum dan Perjanjian Sosial Budaya. '
        'Seluruh modul utama telah dikembangkan dan diuji dengan hasil yang memuaskan.'
    )
    
    doc.add_paragraph()  # Add blank line
    
    # Saran
    para_d2 = doc.add_paragraph()
    para_d2.add_run(
        'Saran: Untuk pengembangan ke depan, disarankan untuk terus melakukan evaluasi dan perbaikan berkelanjutan '
        'berdasarkan masukan dari pengguna. Selain itu, penting untuk memberikan pelatihan kepada pegawai dalam '
        'penggunaan aplikasi agar pemanfaatannya lebih optimal serta memastikan pengguna memahami seluruh fitur yang tersedia.'
    )
    
    doc.add_paragraph()  # Add blank line
    
    # E. Penutup
    heading_e = doc.add_heading('E. Penutup', level=1)
    
    para_e = doc.add_paragraph()
    para_e.add_run(
        'Demikian laporan kegiatan lembur dalam rangka pengembangan Aplikasi Workload Management System. '
        'Kami menyadari bahwa pengembangan aplikasi ini masih akan terus berjalan sesuai dengan kebutuhan '
        'dan masukan dari pengguna. Atas perhatian dan kerja samanya, kami ucapkan terima kasih.'
    )
    
    # Add space
    doc.add_paragraph()
    
    # Add signature area
    doc.add_paragraph('Hormat Kami,', style='Normal')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('.............................', style='Normal')
    doc.add_paragraph('NIP. .............................', style='Normal')
    
    # Save the document
    doc.save('/home/hpsb/workload-app/Laporan/Laporan_Pengembangan_Aplikasi_Workload_Management.docx')
    print('Laporan telah dibuat: /home/hpsb/workload-app/Laporan/Laporan_Pengembangan_Aplikasi_Workload_Management.docx')

if __name__ == '__main__':
    create_report_with_specific_format()