import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.ns import qn
import fitz  # PyMuPDF

def create_comprehensive_report():
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Laporan Pekerjaan - Sistem Workload Management', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    doc.add_paragraph('Laporan ini menyajikan ringkasan pekerjaan yang telah dilakukan dalam pengembangan sistem Workload Management berdasarkan surat perintah kerja dan tampilan aplikasi.', style='Subtitle')
    
    # Add section about the application
    doc.add_heading('Deskripsi Aplikasi', level=1)
    doc.add_paragraph(
        'Sistem Workload Management adalah aplikasi berbasis web yang dirancang untuk mengelola beban kerja '
        'dalam organisasi pemerintahan, khususnya di lingkungan Kementerian Luar Negeri. '
        'Aplikasi ini menyediakan berbagai fungsionalitas untuk memantau dan mengelola tugas-tugas pegawai serta tim.')
    
    # Add UI description
    doc.add_heading('Tampilan dan Fungsionalitas', level=1)
    doc.add_paragraph(
        'Aplikasi memiliki antarmuka yang responsif dengan struktur sebagai berikut:')
    doc.add_paragraph('1. Header dengan logo dan judul "Workload Management System"', style='List Bullet')
    doc.add_paragraph('2. Sidebar navigasi yang dapat diciutkan/diperluas', style='List Bullet')
    doc.add_paragraph('3. Area konten utama dengan berbagai modul', style='List Bullet')
    
    doc.add_paragraph(
        'Fungsionalitas utama aplikasi meliputi:')
    doc.add_paragraph('Dashboard untuk ikhtisar sistem', style='List Bullet')
    doc.add_paragraph('Modul Workload untuk mengelola tugas', style='List Bullet')
    doc.add_paragraph('Modul Team Tasks untuk melihat beban kerja tim', style='List Bullet')
    doc.add_paragraph('Modul Employees untuk manajemen pegawai', style='List Bullet')
    doc.add_paragraph('Modul Calendar untuk penjadwalan', style='List Bullet')
    doc.add_paragraph('Modul Reports untuk laporan dan analitik', style='List Bullet')
    doc.add_paragraph('Modul E-Kinerja untuk dokumentasi kinerja', style='List Bullet')
    doc.add_paragraph('Modul History untuk log audit', style='List Bullet')
    doc.add_paragraph('Modul Settings dan Profile', style='List Bullet')
    
    # Add section about work orders
    doc.add_heading('Surat Perintah Kerja', level=1)
    doc.add_paragraph(
        'Berdasarkan file-file PDF yang tersedia, berikut adalah ringkasan dari surat perintah kerja '
        'yang terkait dengan pengembangan aplikasi ini:')
    
    # Path to PDF files
    pdf_folder = '/home/hpsb/workload-app/Laporan/'
    pdf_files = [
        'Dok-23500-Surat Perintah Kerja Lembur Pembaruan System Aplikasi Administrasion Tools Direktorat Hukum dan Perjanjian Sosial Budaya pada tanggal 8 November 2025-signed (1).pdf',
        'Dok-34545-Surat Perintah Kerja Lembur Pembaruan System Aplikasi Administrasion Tools Direktorat Hukum dan Perjanjian Sosial Budaya pada tanggal 15 November 2025-signed (1).pdf',
        'Dok-65030-Surat Perintah Kerja Lembur Pembaruan System Aplikasi Administrasion Tools Direktorat Hukum dan Perjanjian Sosial Budaya pada tanggal 9 November 2025-signed (1).pdf',
        'Dok-85247-Surat Perintah Kerja Lembur Pembaruan System Aplikasi Administrasion Tools Direktorat Hukum dan Perjanjian Sosial Budaya pada tanggal 16 November 2025-signed (1).pdf'
    ]
    
    for pdf_file in pdf_files:
        pdf_path = os.path.join(pdf_folder, pdf_file)
        
        # Extract information from PDF
        try:
            with fitz.open(pdf_path) as pdf:
                text = ""
                for page_num in range(min(3, len(pdf))):  # Read first 3 pages
                    page = pdf[page_num]
                    text += page.get_text()
            
            # Create a section for each PDF
            doc.add_heading(f'Surat Perintah Kerja: {pdf_file}', level=2)
            
            # Add content extracted from PDF (first few lines as example)
            lines = text.split('\n')
            # Take the first few meaningful lines that likely contain important info
            content_lines = [line.strip() for line in lines if line.strip() and len(line.strip()) > 20]
            preview_text = ' '.join(content_lines[:5])  # First 5 meaningful lines
            
            doc.add_paragraph(f'Isi ringkasan: {preview_text[:500]}...')  # Limit length
            
        except Exception as e:
            doc.add_paragraph(f'Gagal membaca file {pdf_file}: {str(e)}')
    
    # Add a conclusion section
    doc.add_heading('Kesimpulan', level=1)
    doc.add_paragraph(
        'Proyek pengembangan Sistem Workload Management telah melibatkan pembuatan berbagai fitur penting '
        'untuk mendukung manajemen beban kerja di lingkungan Kementerian Luar Negeri. '
        'Surat perintah kerja menunjukkan bahwa pengembangan sistem ini berlangsung secara bertahap '
        'dengan fokus pada pembaruan dan perbaikan aplikasi administrasi.')
    
    doc.add_paragraph(
        'Aplikasi telah menyediakan antarmuka yang lengkap dan fungsional untuk membantu '
        'pegawai dalam mengelola tugas, jadwal, dan dokumentasi kinerja.')
    
    # Save the document
    doc.save('/home/hpsb/workload-app/Laporan/laporan_pekerjaan_sistem_workload_management.docx')
    print('Laporan telah dibuat: /home/hpsb/workload-app/Laporan/laporan_pekerjaan_sistem_workload_management.docx')

if __name__ == '__main__':
    create_comprehensive_report()