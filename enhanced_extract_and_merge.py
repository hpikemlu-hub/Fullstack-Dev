#!/usr/bin/env python3
"""
Script yang lebih canggih untuk mengekstrak informasi dari file PDF surat perintah kerja
dan memasukkannya ke dalam format laporan Word.
"""

import os
import re
from datetime import datetime
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import logging

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def extract_pdf_info(pdf_path):
    """
    Ekstrak informasi penting dari file PDF surat perintah kerja.
    """
    logger.info(f"Mengekstrak informasi dari: {pdf_path}")
    
    doc = fitz.open(pdf_path)
    text = ""
    
    # Ekstrak teks dari semua halaman
    for page in doc:
        text += page.get_text()
    
    doc.close()
    
    # Ekstrak informasi penting menggunakan regex
    info = {}
    
    # Ekstrak nomor surat dari nama file
    filename = Path(pdf_path).stem
    nomor_surat_match = re.search(r'Dok-(\d+)', filename)
    if nomor_surat_match:
        info['nomor_surat'] = nomor_surat_match.group(1)
    else:
        # Coba ekstrak dari isi dokumen
        nomor_surat_doc_match = re.search(r'Nomor[:\s]+([A-Z0-9/]+)', text, re.IGNORECASE)
        if nomor_surat_doc_match:
            info['nomor_surat'] = nomor_surat_doc_match.group(1).strip()
        else:
            info['nomor_surat'] = "Tidak ditemukan"
    
    # Ekstrak tanggal dari nama file
    tanggal_match = re.search(r'pada tanggal (\d+ \w+ \d+)', filename)
    if tanggal_match:
        info['tanggal'] = tanggal_match.group(1)
    else:
        # Coba ekstrak dari isi dokumen
        tanggal_doc_match = re.search(r'tanggal[:\s]+(\d{1,2}\s+\w+\s+\d{4})', text, re.IGNORECASE)
        if tanggal_doc_match:
            info['tanggal'] = tanggal_doc_match.group(1).strip()
        else:
            info['tanggal'] = "Tidak ditemukan"
    
    # Ekstrak deskripsi pekerjaan dari nama file
    deskripsi_match = re.search(r'Surat Perintah Kerja (.+?) pada tanggal', filename)
    if deskripsi_match:
        info['deskripsi_pekerjaan'] = deskripsi_match.group(1).strip()
    else:
        # Coba ekstrak dari isi dokumen
        deskripsi_doc_match = re.search(r'(pekerjaan|kegiatan|pelaksanaan)[:\-\s]+(.*?)(\n|$)', text, re.IGNORECASE | re.DOTALL)
        if deskripsi_doc_match:
            info['deskripsi_pekerjaan'] = deskripsi_doc_match.group(2).strip().split('\n')[0]
        else:
            info['deskripsi_pekerjaan'] = "Tidak ditemukan"
    
    # Coba cari informasi tambahan dari isi PDF
    # Cari tanggal surat dalam isi dokumen
    tanggal_surat_match = re.search(r'tanggal.*?(\d{1,2}\s+\w+\s+\d{4})', text, re.IGNORECASE)
    if tanggal_surat_match:
        info['tanggal_surat_dokumen'] = tanggal_surat_match.group(1)
    
    # Cari informasi tentang kegiatan/pekerjaan lebih lengkap
    kegiatan_patterns = [
        r'(pekerjaan|kegiatan|pelaksanaan)[:\-\s]+(.*?)(\n\s*\n|\n\s*[A-Z][a-z]+\s*:)',
        r'(melaksanakan|melakukan|menugaskan)\s+(.*?)(\n|$)',
        r'(uraian|deskripsi)\s+(pekerjaan|kegiatan)[:\-\s]+(.*?)(\n\s*\n|\n\s*[A-Z][a-z]+\s*:)',
    ]
    
    for pattern in kegiatan_patterns:
        kegiatan_matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
        if kegiatan_matches:
            for match in kegiatan_matches:
                kegiatan = match[-1].strip()
                if len(kegiatan) > 10:  # Pastikan cukup panjang untuk menjadi deskripsi
                    info['deskripsi_lengkap'] = kegiatan.replace('\n', ' ').strip()
                    break
            if 'deskripsi_lengkap' in info:
                break
    
    # Cari nama-nama orang yang terlibat
    nama_patterns = [
        r'(Kepala Bidang|Kepala Subbagian|Pejabat|Penanggung Jawab|Menyetujui|Mengetahui|Yang Bertanda Tangan|Atasan Langsung)[:\-\s]+([A-Z\s\.]+?)(\n|$)',
        r'(dengan ini menugaskan|dengan ini memerintahkan)\s+([A-Z\s\.]+?)(untuk|melaksanakan|menyelesaikan)',
    ]
    
    penandatangan = []
    for pattern in nama_patterns:
        nama_matches = re.findall(pattern, text, re.IGNORECASE)
        if nama_matches:
            for match in nama_matches:
                nama = match[-1].strip()
                if len(nama) > 2 and len(nama.split()) >= 2:  # Pastikan nama valid
                    penandatangan.append(nama)
    
    if penandatangan:
        info['penandatangan'] = penandatangan
    
    # Cari durasi atau waktu pelaksanaan
    durasi_match = re.search(r'(tanggal|periode|waktu pelaksanaan)[:\-\s]+(.*?)(\n|$)', text, re.IGNORECASE)
    if durasi_match:
        info['durasi_pelaksanaan'] = durasi_match.group(2).strip()
    
    # Cari lokasi pelaksanaan
    lokasi_match = re.search(r'(lokasi|tempat pelaksanaan|di|pada)[:\-\s]+(.*?)(\n|$)', text, re.IGNORECASE)
    if lokasi_match:
        info['lokasi_pelaksanaan'] = lokasi_match.group(2).strip()
    
    logger.info(f"Ekstraksi selesai untuk {info['nomor_surat']}")
    return info

def insert_info_to_word(word_path, pdf_info_list, output_path):
    """
    Masukkan informasi dari PDF ke dalam file Word.
    """
    logger.info(f"Memasukkan informasi ke dalam file Word: {word_path}")
    
    # Buka dokumen Word
    doc = Document(word_path)
    
    # Tambahkan judul bagian untuk informasi surat perintah kerja
    heading = doc.add_heading('Daftar Surat Perintah Kerja', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Urutkan informasi berdasarkan tanggal
    sorted_info = sorted(pdf_info_list, key=lambda x: x['tanggal'])
    
    # Tambahkan tabel untuk menampilkan informasi surat perintah kerja
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Header tabel
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nomor Surat'
    hdr_cells[1].text = 'Tanggal'
    hdr_cells[2].text = 'Deskripsi Pekerjaan'
    hdr_cells[3].text = 'Durasi Pelaksanaan'
    hdr_cells[4].text = 'Keterangan Tambahan'
    
    # Tambahkan baris untuk setiap surat perintah kerja
    for info in sorted_info:
        row_cells = table.add_row().cells
        row_cells[0].text = f"Dok-{info['nomor_surat']}" if info['nomor_surat'].isdigit() else info['nomor_surat']
        row_cells[1].text = info['tanggal']
        row_cells[2].text = info['deskripsi_pekerjaan']
        
        # Tambahkan durasi pelaksanaan jika ada
        if 'durasi_pelaksanaan' in info:
            row_cells[3].text = info['durasi_pelaksanaan']
        else:
            row_cells[3].text = "Tidak disebutkan"
        
        # Tambahkan keterangan tambahan jika ada
        keterangan = []
        if 'tanggal_surat_dokumen' in info and info['tanggal_surat_dokumen'] != info['tanggal']:
            keterangan.append(f"Tanggal Dokumen: {info['tanggal_surat_dokumen']}")
        if 'deskripsi_lengkap' in info:
            keterangan.append(f"Deskripsi: {info['deskripsi_lengkap'][:100]}...")  # Ambil 100 karakter pertama
        if 'penandatangan' in info:
            keterangan.append(f"Penandatangan: {', '.join(info['penandatangan'][:2])}")  # Ambil 2 pertama
        if 'lokasi_pelaksanaan' in info:
            keterangan.append(f"Lokasi: {info['lokasi_pelaksanaan']}")
        
        row_cells[4].text = '\n'.join(keterangan) if keterangan else "Tidak ada keterangan tambahan"
    
    # Tambahkan spasi
    doc.add_paragraph()
    
    # Tambahkan detail untuk setiap surat perintah kerja
    for i, info in enumerate(sorted_info):
        doc.add_heading(f'Surat Perintah Kerja {i+1}: Dok-{info["nomor_surat"] if info["nomor_surat"].isdigit() else info["nomor_surat"]}', level=2)
        
        # Tambahkan detail dalam bentuk paragraf
        doc.add_paragraph(f"Nomor Surat: Dok-{info['nomor_surat'] if info['nomor_surat'].isdigit() else info['nomor_surat']}")
        doc.add_paragraph(f"Tanggal: {info['tanggal']}")
        doc.add_paragraph(f"Deskripsi Pekerjaan: {info['deskripsi_pekerjaan']}")
        
        if 'tanggal_surat_dokumen' in info and info['tanggal_surat_dokumen'] != info['tanggal']:
            doc.add_paragraph(f"Tanggal Dokumen: {info['tanggal_surat_dokumen']}")
        
        if 'deskripsi_lengkap' in info:
            doc.add_paragraph(f"Deskripsi Lengkap: {info['deskripsi_lengkap']}")
        
        if 'penandatangan' in info:
            doc.add_paragraph(f"Penandatangan: {', '.join(info['penandatangan'])}")
        
        if 'durasi_pelaksanaan' in info:
            doc.add_paragraph(f"Durasi Pelaksanaan: {info['durasi_pelaksanaan']}")
        
        if 'lokasi_pelaksanaan' in info:
            doc.add_paragraph(f"Lokasi Pelaksanaan: {info['lokasi_pelaksanaan']}")
        
        doc.add_paragraph()  # Spasi antar surat
    
    # Simpan dokumen yang telah dimodifikasi
    doc.save(output_path)
    logger.info(f"File Word berhasil disimpan sebagai: {output_path}")

def main():
    """
    Fungsi utama untuk mengekstrak informasi dari PDF dan memasukkan ke Word.
    """
    # Direktori tempat file berada
    directory = Path("/home/hpsb/workload-app/Laporan")
    
    # Ambil semua file PDF di direktori Laporan yang merupakan surat perintah kerja
    pdf_files = list(directory.glob("Dok-*.pdf"))
    
    # File Word asli dan file output
    word_file = directory / "Laporan Pengembangan Aplikasi 18 - 19, 28 - 29 Oktober 2025.docx"
    output_file = directory / "Laporan_Pengembangan_Aplikasi_dengan_SPK_enhanced.docx"
    
    # Pastikan file-file ada
    if not pdf_files:
        logger.error("Tidak ditemukan file PDF surat perintah kerja")
        return
    
    if not word_file.exists():
        logger.error(f"File Word tidak ditemukan: {word_file}")
        return
    
    # Ekstrak informasi dari semua file PDF
    pdf_info_list = []
    for pdf_file in pdf_files:
        try:
            info = extract_pdf_info(str(pdf_file))
            pdf_info_list.append(info)
        except Exception as e:
            logger.error(f"Error saat mengekstrak informasi dari {pdf_file}: {e}")
    
    # Masukkan informasi ke dalam file Word
    try:
        insert_info_to_word(str(word_file), pdf_info_list, str(output_file))
        logger.info("Proses ekstraksi dan penggabungan selesai.")
        print(f"File laporan telah dibuat: {output_file}")
    except Exception as e:
        logger.error(f"Error saat memasukkan informasi ke file Word: {e}")

if __name__ == "__main__":
    main()