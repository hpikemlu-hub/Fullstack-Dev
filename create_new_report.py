#!/usr/bin/env python3
"""
Script untuk membuat laporan baru sesuai dengan format dari template
'Laporan Pengembangan Aplikasi 18 - 19, 28 - 29 Oktober 2025.docx'
"""

import os
import re
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor
import logging

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def extract_pdf_info(pdf_path):
    """
    Ekstrak informasi penting dari file PDF surat perintah kerja.
    """
    logger.info(f"Mengekstrak informasi dari: {pdf_path}")

    doc = fitz.open(pdf_path)
    text = ""

    # Ekstrak teks dari semua halaman
    for page in doc:
        text += page.get_text()

    doc.close()

    # Ekstrak informasi penting menggunakan regex
    info = {}

    # Ekstrak nomor surat dari nama file
    filename = Path(pdf_path).stem
    nomor_surat_match = re.search(r'Dok-(\d+)', filename)
    if nomor_surat_match:
        info['nomor_surat'] = nomor_surat_match.group(1)
    else:
        # Coba ekstrak dari isi dokumen
        nomor_surat_doc_match = re.search(r'Nomor[:\s]+([A-Z0-9/]+)', text, re.IGNORECASE)
        if nomor_surat_doc_match:
            info['nomor_surat'] = nomor_surat_doc_match.group(1).strip()
        else:
            info['nomor_surat'] = "Tidak ditemukan"

    # Ekstrak tanggal dari nama file
    tanggal_match = re.search(r'pada tanggal (\d+ \w+ \d+)', filename)
    if tanggal_match:
        info['tanggal'] = tanggal_match.group(1)
    else:
        # Coba ekstrak dari isi dokumen
        tanggal_doc_match = re.search(r'tanggal[:\s]+(\d{1,2}\s+\w+\s+\d{4})', text, re.IGNORECASE)
        if tanggal_doc_match:
            info['tanggal'] = tanggal_doc_match.group(1).strip()
        else:
            info['tanggal'] = "Tidak ditemukan"

    # Ekstrak deskripsi pekerjaan dari nama file
    deskripsi_match = re.search(r'Surat Perintah Kerja (.+?) pada tanggal', filename)
    if deskripsi_match:
        info['deskripsi_pekerjaan'] = deskripsi_match.group(1).strip()
    else:
        # Coba ekstrak dari isi dokumen
        deskripsi_doc_match = re.search(r'(pekerjaan|kegiatan|pelaksanaan)[:\-\s]+(.*?)(\n|$)', text, re.IGNORECASE | re.DOTALL)
        if deskripsi_doc_match:
            info['deskripsi_pekerjaan'] = deskripsi_doc_match.group(2).strip().split('\n')[0]
        else:
            info['deskripsi_pekerjaan'] = "Tidak ditemukan"

    # Coba cari informasi tambahan dari isi PDF
    # Cari tanggal surat dalam isi dokumen
    tanggal_surat_match = re.search(r'tanggal.*?(\d{1,2}\s+\w+\s+\d{4})', text, re.IGNORECASE)
    if tanggal_surat_match:
        info['tanggal_surat_dokumen'] = tanggal_surat_match.group(1)

    # Cari informasi tentang kegiatan/pekerjaan lebih lengkap
    kegiatan_patterns = [
        r'(pekerjaan|kegiatan|pelaksanaan)[:\-\s]+(.*?)(\n\s*\n|\n\s*[A-Z][a-z]+\s*:)',
        r'(melaksanakan|melakukan|menugaskan)\s+(.*?)(\n|$)',
        r'(uraian|deskripsi)\s+(pekerjaan|kegiatan)[:\-\s]+(.*?)(\n\s*\n|\n\s*[A-Z][a-z]+\s*:)',
    ]

    for pattern in kegiatan_patterns:
        kegiatan_matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
        if kegiatan_matches:
            for match in kegiatan_matches:
                kegiatan = match[-1].strip()
                if len(kegiatan) > 10:  # Pastikan cukup panjang untuk menjadi deskripsi
                    info['deskripsi_lengkap'] = kegiatan.replace('\n', ' ').strip()
                    break
            if 'deskripsi_lengkap' in info:
                break

    # Cari nama-nama orang yang terlibat
    nama_patterns = [
        r'(Kepala Bidang|Kepala Subbagian|Pejabat|Penanggung Jawab|Menyetujui|Mengetahui|Yang Bertanda Tangan|Atasan Langsung)[:\-\s]+([A-Z\s\.]+?)(\n|$)',
        r'(dengan ini menugaskan|dengan ini memerintahkan)\s+([A-Z\s\.]+?)(untuk|melaksanakan|menyelesaikan)',
    ]

    penandatangan = []
    for pattern in nama_patterns:
        nama_matches = re.findall(pattern, text, re.IGNORECASE)
        if nama_matches:
            for match in nama_matches:
                nama = match[-1].strip()
                if len(nama) > 2 and len(nama.split()) >= 2:  # Pastikan nama valid
                    penandatangan.append(nama)

    if penandatangan:
        info['penandatangan'] = penandatangan

    # Cari durasi atau waktu pelaksanaan
    durasi_match = re.search(r'(tanggal|periode|waktu pelaksanaan)[:\-\s]+(.*?)(\n|$)', text, re.IGNORECASE)
    if durasi_match:
        info['durasi_pelaksanaan'] = durasi_match.group(2).strip()

    # Cari lokasi pelaksanaan
    lokasi_match = re.search(r'(lokasi|tempat pelaksanaan|di|pada)[:\-\s]+(.*?)(\n|$)', text, re.IGNORECASE)
    if lokasi_match:
        info['lokasi_pelaksanaan'] = lokasi_match.group(2).strip()

    logger.info(f"Ekstraksi selesai untuk {info['nomor_surat']}")
    return info

def create_new_report():
    """
    Create a new report following the template format from 'Laporan Pengembangan Aplikasi 18 - 19, 28 - 29 Oktober 2025.docx'
    """
    logger.info("Membuat laporan baru sesuai format template...")

    # Create a new document
    doc = Document()

    # Add title following the template format
    title = doc.add_heading('Laporan Pengembangan Aplikasi Workload Management', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Format the title with bold and larger font
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    # Add subtitle
    doc.add_paragraph('Laporan ini menyajikan ringkasan pekerjaan yang telah dilakukan dalam pengembangan sistem Workload Management berdasarkan surat perintah kerja dan tampilan aplikasi.', style='Subtitle')

    # Add section about the application following template structure
    doc.add_heading('Deskripsi Aplikasi', level=1)
    
    app_description = doc.add_paragraph()
    app_description.add_run('Sistem Workload Management adalah aplikasi berbasis web yang dirancang untuk mengelola beban kerja dalam organisasi pemerintahan, khususnya di lingkungan Kementerian Luar Negeri. Aplikasi ini menyediakan berbagai fungsionalitas untuk memantau dan mengelola tugas-tugas pegawai serta tim.')
    
    # Add UI description
    doc.add_heading('Tampilan dan Fungsionalitas', level=1)
    
    ui_description = doc.add_paragraph()
    ui_description.add_run('Aplikasi memiliki antarmuka yang responsif dengan struktur sebagai berikut:')
    
    doc.add_paragraph('1. Header dengan logo dan judul "Workload Management System"', style='List Number')
    doc.add_paragraph('2. Sidebar navigasi yang dapat diciutkan/diperluas', style='List Number')
    doc.add_paragraph('3. Area konten utama dengan berbagai modul', style='List Number')

    doc.add_paragraph('Fungsionalitas utama aplikasi meliputi:', style='List Paragraph')
    doc.add_paragraph('- Dashboard untuk ikhtisar sistem', style='List Bullet')
    doc.add_paragraph('- Modul Workload untuk mengelola tugas', style='List Bullet')
    doc.add_paragraph('- Modul Team Tasks untuk melihat beban kerja tim', style='List Bullet')
    doc.add_paragraph('- Modul Employees untuk manajemen pegawai', style='List Bullet')
    doc.add_paragraph('- Modul Calendar untuk penjadwalan', style='List Bullet')
    doc.add_paragraph('- Modul Reports untuk laporan dan analitik', style='List Bullet')
    doc.add_paragraph('- Modul E-Kinerja untuk dokumentasi kinerja', style='List Bullet')
    doc.add_paragraph('- Modul History untuk log audit', style='List Bullet')
    doc.add_paragraph('- Modul Settings dan Profile', style='List Bullet')

    # Add screenshots section
    doc.add_heading('Tampilan Aplikasi', level=1)
    
    screenshots_description = doc.add_paragraph()
    screenshots_description.add_run('Berikut adalah beberapa screenshot tampilan aplikasi Workload Management System:')
    
    # Add screenshots from the Gambar directory
    gambar_dir = Path("/home/hpsb/workload-app/Laporan/Gambar")
    screenshot_files = list(gambar_dir.glob("*.png"))
    
    if screenshot_files:
        for i, screenshot_path in enumerate(screenshot_files[:6]):  # Limit to first 6 screenshots to avoid making the doc too large
            try:
                # Add screenshot to document
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(screenshot_path), width=Inches(6.5))  # Set width to fit page
                
                # Add caption
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(f'Gambar {i+1}: Tampilan {screenshot_path.name}')
                caption_run.font.size = Pt(10)
                
            except Exception as e:
                logger.error(f"Gagal menambahkan screenshot {screenshot_path}: {e}")
                doc.add_paragraph(f'Gambar {i+1}: Gagal memuat {screenshot_path.name} ({str(e)})')
    else:
        doc.add_paragraph('Tidak ditemukan screenshot dalam direktori Gambar.')

    # Add section about work orders
    doc.add_heading('Surat Perintah Kerja', level=1)
    
    work_orders_description = doc.add_paragraph()
    work_orders_description.add_run('Berdasarkan file-file PDF yang tersedia, berikut adalah ringkasan dari surat perintah kerja yang terkait dengan pengembangan aplikasi ini:')
    
    # Path to PDF files
    pdf_folder = Path('/home/hpsb/workload-app/Laporan/')
    pdf_files = list(pdf_folder.glob("Dok-*.pdf"))

    if pdf_files:
        # Ekstrak informasi dari semua file PDF
        pdf_info_list = []
        for pdf_file in pdf_files:
            try:
                info = extract_pdf_info(str(pdf_file))
                pdf_info_list.append(info)
            except Exception as e:
                logger.error(f"Error saat mengekstrak informasi dari {pdf_file}: {e}")
                doc.add_paragraph(f'Gagal membaca file {pdf_file.name}: {str(e)}')

        # Urutkan informasi berdasarkan tanggal
        sorted_info = sorted(pdf_info_list, key=lambda x: x['tanggal'])

        # Tambahkan tabel untuk menampilkan informasi surat perintah kerja
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Header tabel
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nomor Surat'
        hdr_cells[1].text = 'Tanggal'
        hdr_cells[2].text = 'Deskripsi Pekerjaan'
        hdr_cells[3].text = 'Durasi Pelaksanaan'
        hdr_cells[4].text = 'Keterangan Tambahan'

        # Tambahkan baris untuk setiap surat perintah kerja
        for info in sorted_info:
            row_cells = table.add_row().cells
            row_cells[0].text = f"Dok-{info['nomor_surat']}" if info['nomor_surat'].isdigit() else info['nomor_surat']
            row_cells[1].text = info['tanggal']
            row_cells[2].text = info['deskripsi_pekerjaan']

            # Tambahkan durasi pelaksanaan jika ada
            if 'durasi_pelaksanaan' in info:
                row_cells[3].text = info['durasi_pelaksanaan']
            else:
                row_cells[3].text = "Tidak disebutkan"

            # Tambahkan keterangan tambahan jika ada
            keterangan = []
            if 'tanggal_surat_dokumen' in info and info['tanggal_surat_dokumen'] != info['tanggal']:
                keterangan.append(f"Tanggal Dokumen: {info['tanggal_surat_dokumen']}")
            if 'deskripsi_lengkap' in info:
                keterangan.append(f"Deskripsi: {info['deskripsi_lengkap'][:100]}...")  # Ambil 100 karakter pertama
            if 'penandatangan' in info:
                keterangan.append(f"Penandatangan: {', '.join(info['penandatangan'][:2])}")  # Ambil 2 pertama
            if 'lokasi_pelaksanaan' in info:
                keterangan.append(f"Lokasi: {info['lokasi_pelaksanaan']}")

            row_cells[4].text = '\n'.join(keterangan) if keterangan else "Tidak ada keterangan tambahan"

        # Tambahkan spasi
        doc.add_paragraph()

        # Tambahkan detail untuk setiap surat perintah kerja
        for i, info in enumerate(sorted_info):
            doc.add_heading(f'Surat Perintah Kerja {i+1}: Dok-{info["nomor_surat"] if info["nomor_surat"].isdigit() else info["nomor_surat"]}', level=2)

            # Tambahkan detail dalam bentuk paragraf
            doc.add_paragraph(f"Nomor Surat: Dok-{info['nomor_surat'] if info['nomor_surat'].isdigit() else info['nomor_surat']}")
            doc.add_paragraph(f"Tanggal: {info['tanggal']}")
            doc.add_paragraph(f"Deskripsi Pekerjaan: {info['deskripsi_pekerjaan']}")

            if 'tanggal_surat_dokumen' in info and info['tanggal_surat_dokumen'] != info['tanggal']:
                doc.add_paragraph(f"Tanggal Dokumen: {info['tanggal_surat_dokumen']}")

            if 'deskripsi_lengkap' in info:
                doc.add_paragraph(f"Deskripsi Lengkap: {info['deskripsi_lengkap']}")

            if 'penandatangan' in info:
                doc.add_paragraph(f"Penandatangan: {', '.join(info['penandatangan'])}")

            if 'durasi_pelaksanaan' in info:
                doc.add_paragraph(f"Durasi Pelaksanaan: {info['durasi_pelaksanaan']}")

            if 'lokasi_pelaksanaan' in info:
                doc.add_paragraph(f"Lokasi Pelaksanaan: {info['lokasi_pelaksanaan']}")

            doc.add_paragraph()  # Spasi antar surat
    else:
        doc.add_paragraph('Tidak ditemukan file surat perintah kerja (PDF) dalam direktori.')

    # Add a conclusion section
    doc.add_heading('Kesimpulan', level=1)
    
    conclusion1 = doc.add_paragraph()
    conclusion1.add_run('Proyek pengembangan Sistem Workload Management telah melibatkan pembuatan berbagai fitur penting untuk mendukung manajemen beban kerja di lingkungan Kementerian Luar Negeri. Surat perintah kerja menunjukkan bahwa pengembangan sistem ini berlangsung secara bertahap dengan fokus pada pembaruan dan perbaikan aplikasi administrasi.')
    
    doc.add_paragraph()
    
    conclusion2 = doc.add_paragraph()
    conclusion2.add_run('Aplikasi telah menyediakan antarmuka yang lengkap dan fungsional untuk membantu pegawai dalam mengelola tugas, jadwal, dan dokumentasi kinerja. Integrasi antara surat perintah kerja dan implementasi teknis menunjukkan pendekatan yang sistematis dalam pengembangan aplikasi ini.')
    
    # Save the document
    output_path = '/home/hpsb/workload-app/Laporan/Laporan_Pengembangan_Aplikasi_Baru.docx'
    doc.save(output_path)
    print(f'Laporan baru telah dibuat: {output_path}')
    logger.info(f'Laporan baru telah dibuat: {output_path}')

if __name__ == '__main__':
    create_new_report()